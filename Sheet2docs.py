#made by Yarko Bahamonde
#distribution outside of my github page is prohibited
#can be used in non-profit projects
#the complete or partial use of the code for other projects needs to be credited

import os
from docx import Document
from openpyxl import load_workbook

# Find the sheet file
sheet_path = None
for file in os.listdir('.'):
    if file.endswith('.xlsx'):
        sheet_path = file
        break

if not sheet_path:
    raise FileNotFoundError("No sheet file found.")


# Load Excel data
wb = load_workbook(sheet_path)
ws = wb.active

# Find the template file
template_path = None
for file in os.listdir('.'):
    if file.endswith('.docx') and file.startswith('template'):
        template_path = file
        break

if not template_path:
    raise FileNotFoundError("No template file found with 'template' in the name.")


# Read the first row
first_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))

#print("First row data:", first_row)

headers = list(filter(lambda x: x is not None, first_row))

# Iterate through each row in Excel (skip the header row)
for row in ws.iter_rows(min_row=2, values_only=True):  # min_row=2 skips the header

    #conuter files saved
    files_saved = 0

    # Load the Word template
    doc = Document(template_path)

    # Replace placeholders with actual data
    for paragraph in doc.paragraphs:
        for header in headers:
            if f'<<{header}>>' in paragraph.text:
                indexHeader = headers.index(header)
                if row[indexHeader] is None or row[indexHeader] == '':
                    paragraph.text = paragraph.text.replace(f'<<{header}>>', ' ')

                if "img" in header:
                    paragraph.text = paragraph.text.replace(f'<<{header}>>', '')
                    run = paragraph.add_run()
                    run.add_picture(row[indexHeader])
                    #doc.add_picture(row[indexHeader])
                else:
                    paragraph.text = paragraph.text.replace(f'<<{header}>>', str(row[indexHeader]))

    # Save the document with a custom name
    output_path = template_path.replace("template", "").strip().replace(" ", "_")

    for header in headers:
        #print(header)
        if ('--'+header+'--') in output_path:
            #print(row[headers.index(header)])
            output_path = output_path.replace('--'+header+'--', str(row[headers.index(header)]))
            
    # Ensure the output path is unique
    if os.path.exists(output_path):
        counter = 1
        output_path = output_path.rsplit('.', 1)[0] + f'({counter}).docx'
        while os.path.exists(output_path):
            # Increment the counter to create a unique filename
            counter += 1
            output_path = output_path.rsplit('(', 1)[0] + f'({counter}).docx'
    #print(output_path)
    doc.save(output_path)
    print(f'Saved: {output_path}')
    files_saved += 1

input(files_saved + " documents generated successfully!")

